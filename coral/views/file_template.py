# coding: utf-8
"""
ARCHES - a program developed to inventory and manage immovable cultural heritage.
Copyright (C) 2013 J. Paul Getty Trust and World Monuments Fund

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU Affero General Public License as
published by the Free Software Foundation, either version 3 of the
License, or (at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
GNU Affero General Public License for more details.

You should have received a copy of the GNU Affero General Public License
along with this program. If not, see <http://www.gnu.org/licenses/>.
"""

import json
import os
import re
import uuid
from datetime import datetime
from typing import Dict, Any, List, NewType
from collections.abc import MutableMapping
import docx
import textwrap
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.parser import OxmlElement
from html.parser import HTMLParser
from html.entities import name2codepoint
from django.core.files.uploadedfile import UploadedFile
from django.http import HttpRequest, HttpResponseNotFound
from django.utils.translation import gettext as _
from django.views.generic import View
from arches.app.datatypes.datatypes import DataTypeFactory
from arches.app.models import models
from arches.app.models.concept import ConceptValue
from arches.app.models.resource import Resource
from arches.app.models.system_settings import settings
from arches.app.models.tile import Tile
from arches.app.utils.response import JSONResponse
from arches.app.views.tile import TileData
import arches_orm
from arches_orm.wkrm import get_well_known_resource_model_by_graph_id
import pytz
from django.core.files.storage import  default_storage
from coral.views.pdf_extract import PdfExtract
import os
class FileTemplateView(View):
    def __init__(self):
        self.doc = None
        self.resource = None

    def get(self, request):
        parenttile_id = request.GET.get("parenttile_id")
        parent_tile = Tile.objects.get(tileid=parenttile_id)
        letter_tiles = Tile.objects.filter(parenttile=parent_tile)
        file_list_node_id = "96f8830a-8490-11ea-9aba-f875a44e0e11"  # Digital Object
        url = None
        for tile in letter_tiles:
            if url is not None:
                break
            for data_obj in tile.data[file_list_node_id]:
                if data_obj["status"] == "uploaded":
                    url = data_obj["url"]
                    break

        if url is not None:
            return JSONResponse({"msg": "success", "download": url})
        return HttpResponseNotFound("No letters tile matching query by parent tile")

    def post(self, request):
        data = json.loads(request.body.decode("utf-8"))
        template_id = request.POST.get("template_id", data.get("template_id", None))
        config = request.POST.get("config", data.get("config", {}))
        extract_pdf = config.get("extract_pdf", False)
        parenttile_id = request.POST.get("parenttile_id")
        resourceinstance_id = request.POST.get(
            "resourceinstance_id", data.get("resourceinstance_id", None)
        )
        transaction_id = request.POST.get("transaction_id", uuid.uuid1())
        self.resource = Resource.objects.get(resourceinstanceid=resourceinstance_id)
        self.resource.load_tiles()
        self.user = request.user
        config["user"] = self.user

        if (
            os.path.exists(os.path.join(settings.APP_ROOT, "uploadedfiles", "docx"))
            is False
        ):
            os.mkdir(os.path.join(settings.APP_ROOT, "uploadedfiles", "docx"))

        fs = default_storage
        template_dict = self.get_template_path(template_id)
        template_path = None
        filesystem_class = default_storage.__class__.__name__
        if filesystem_class == 'S3Boto3Storage':
            template_path = os.path.join(
                "docx", template_dict["filename"]
            )
        elif filesystem_class == 'FileSystemStorage':
            template_path = os.path.join(
                settings.APP_ROOT, "docx", template_dict["filename"]
            )
        try:
            self.doc = Document(fs.open(template_path))
        except:
            return HttpResponseNotFound("No Template Found")
        
        if extract_pdf:
             config['extract_pdf_text'] = []
             pdf_extract = PdfExtract()
             files = data.get('files')
             for file in files:
                filename = file["name"].replace(" ", "_")
                if filesystem_class == 'S3Boto3Storage':
                    file_path = os.path.join(
                        "uploadedfiles", filename
                    )
                elif filesystem_class == 'FileSystemStorage':
                    file_path = os.path.join(
                        settings.APP_ROOT, "uploadedfiles", filename
                    )
                try:
                    file = fs.open(file_path)
                    text = pdf_extract.extract_text(file.read())
                    config['extract_pdf_text'].append(
                        { 
                              'file': file, 
                              'text': text 
                        }
                    )
                except:
                    return HttpResponseNotFound(f"No file found")

        self.edit_letter(self.resource, template_dict["provider"], config)

        timezone = pytz.timezone("Europe/London") 
        current_datetime = datetime.now(timezone)
        # Date and time as "DD-MM-YYYY-HH-MM"
        formatted_datetime = current_datetime.strftime("%d-%m-%Y-%H-%M")

        new_file_name = formatted_datetime + "_" + template_dict["filename"]
        new_file_path = os.path.join(
            settings.APP_ROOT, "uploadedfiles/docx", new_file_name
        )

        new_req = HttpRequest()
        new_req.method = "POST"
        new_req.user = request.user
        new_req.POST["data"] = None
        host = request.get_host()

        self.doc.save(new_file_path)
        saved_file = open(new_file_path, "rb")
        stat = os.stat(new_file_path)
        file_data = UploadedFile(saved_file)
        file_list_node_id = "96f8830a-8490-11ea-9aba-f875a44e0e11"  # Digital Object

        tile = json.dumps(
            {
                "tileid": None,
                "data": {
                    file_list_node_id: [
                        {
                            "name": new_file_name,
                            "accepted": True,
                            "height": 0,
                            "lastModified": stat.st_mtime,
                            "size": stat.st_size,
                            "status": "queued",
                            "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            "width": 0,
                            "url": None,
                            "file_id": None,
                            "index": 0,
                            "content": "blob:" + host + "/{0}".format(uuid.uuid4()),
                        }
                    ]
                },
                "nodegroup_id": "7db68c6c-8490-11ea-a543-f875a44e0e11",
                "parenttile_id": None,
                "resourceinstance_id": "",
                "sortorder": 0,
                "tiles": {},
            }
        )

        new_req = HttpRequest()
        new_req.method = "POST"
        new_req.user = request.user
        new_req.POST["data"] = tile
        new_req.POST["transaction_id"] = transaction_id
        new_req.FILES["file-list_" + file_list_node_id] = file_data
        new_tile = TileData()
        new_tile.action = "update_tile"

        response = TileData.post(new_tile, new_req)
        if response.status_code == 200:
            tile = json.loads(response.content)
            return JSONResponse({"tile": tile, "status": "success"})

        return HttpResponseNotFound(response.status_code)

    def save(self, tile, request, context=None):
        raise NotImplementedError


    def post_save(self, *args, **kwargs):
        raise NotImplementedError

    def delete(self, tile, request):
        raise NotImplementedError

    def on_import(self, tile):
        raise NotImplementedError


    def after_function_save(self, tile, request):
        raise NotImplementedError

    def get_template_path(self, template_id):
        template_dict = {  # keys are valueids from "Letters" concept list; values are known file names
             "a01fb4b1-b443-f59b-e086-dd647de66fde": {
                "filename": "forestry-consult-response-template.docx",
                "provider": GenericTemplateProvider
            },
            "f9bb9268-810a-ceff-754d-a18e92c9d8a1": {
                "filename": "daera-consultation-response-template.docx",
                "provider": GenericTemplateProvider
            },
            "c5562218-ed5f-4468-9efd-656e3337badd": {
                "filename": "daera-ail-consultation-template.docx",
                "provider": GenericTemplateProvider
            },
            "7f1e7061-8bb0-4338-9342-118f1e9214aa": {
                "filename": "smc-addendum-template.docx",
                "provider": GenericTemplateProvider
            },
            "e14bd058-e9f2-48f8-8ef5-337310c3420f": {
                "filename": "smc-provisional-template.docx",
                "provider": GenericTemplateProvider
            },
            "8d605e5c-d0da-4b72-9ce3-2f7dac3381d1": {
                "filename": "smc-refusal-template.docx",
                "provider": GenericTemplateProvider
            },
            "64865b40-1ab8-d59e-aacf-668b077f0e9a" : {
                "filename": "licence-covering-letter.docx",
                "provider": GenericTemplateProvider
            },
            "e3a1469c-ae0e-7a59-2e25-09c036647e89": {
                "filename": "final-report-letter.docx",
                "provider": GenericTemplateProvider,
            },
            "a346c0f8-7be6-4a7c-da2c-aef42b02ee26": {
                "filename": "extra-name-on-letter.docx",
                "provider": GenericTemplateProvider,
            },
            "28e3fbdd-8dc8-4fb9-e378-91df6776e065": {
                "filename": "extension-of-licence-letter.docx",
                "provider": GenericTemplateProvider,
            },
            "bb6e23cf-0331-ea39-7164-a3af1607a958": {
                "filename": "advance-listing-letter-to-owner-occupier.docx",
                "provider": GenericTemplateProvider,
            },
            "95456a6b-30e0-b727-91f9-ccc0acfcac87": {
                "filename": "advance-listing-letter-to-council.docx",
                "provider": GenericTemplateProvider,
            },
            "28faad65-2403-b0b9-3ba8-bbbc99efd175": {
                "filename": "overdue-excavations-report-letter.docx",
                "provider": GenericTemplateProvider,
            },
            "f08c5baf-4434-ac59-a02c-e692ae79d455": {
                "filename": "transfer-of-licence-letter.docx",
                "provider": GenericTemplateProvider,
            },
            "planning-pdf-merger": {
                 "filename": "planning-response-combined.docx",
                 "provider": GenericTemplateProvider
            }
        }
        for key, value in list(template_dict.items()):
            if key == template_id:
                return value
        return { "filename": template_id, "provider": GenericTemplateProvider}

    def edit_letter(self, resource, provider, config):
        # include = []

        # if len(self.doc.paragraphs) > 0:
        #     for paragraph in self.doc.paragraphs:
        #          if re.search('<.*>', paragraph.text):
        #             include += re.findall('<([^<]*)>', paragraph.text)

        # if len(self.doc.tables) > 0:
        #     for table in self.doc.tables:
        #         for row in table.rows:
        #             for cell in row.cells:
        #                  if re.search('<.*>', cell.text):
        #                     include += re.findall('<([^<]*)>', cell.text)

        # config["include"] = include
        mapping_dict = provider(resource).get_mapping(config)

        if config.get("extract_pdf", False):
             data = config.get("extract_pdf_text", [])
             mapping_dict["pdf_data"] = data

        self.apply_mapping(mapping_dict)

    def apply_mapping(self, mapping):
        htmlTags = re.compile(r"<(?:\"[^\"]*\"['\"]*|'[^']*'['\"]*|[^'\">])+>")
        for key in mapping:
            if (isinstance(mapping[key], (arches_orm.view_models.concepts.EmptyConceptValueViewModel))):
                mapping[key] = None
            if (isinstance(mapping[key], (bool))):
                mapping[key] = str(mapping[key])
            if (isinstance(mapping[key], (list))):
                if all(isinstance(item, str) for item in mapping[key]):
                    mapping[key] = "\n \n".join(mapping[key])
            html = False
            if htmlTags.search(str(mapping.get(key, ""))):
                html = True
            self.replace_string(self.doc, key, mapping[key], html)

    def replace_string(self, document, key, v, html=False):
        # Note that the intent here is to preserve how things are styled in the docx
        # easiest way is to iterate through p.runs, not as fast as iterating through parent.paragraphs
        # advantage of the former is that replacing run.text preserves styling, replacing p.text does not

        def parse_html_to_docx(p, k, v):
            style = p.style
            if k in p.text:
                p.clear()
                document_html_parser = DocumentHTMLParser(p, document)
                document_html_parser.insert_into_paragraph_and_feed(v)

        def replace_in_runs(p_list, k, v):
            for paragraph in p_list:
                if html is True:
                    parse_html_to_docx(paragraph, k, v)
                for i, run in enumerate(paragraph.runs):
                    if k in run.text:  # now check if html
                        run_style = run.style
                        run.text = run.text.replace(k, v)
                    elif (
                        i == (len(paragraph.runs) - 1) and k in paragraph.text
                    ):  # backstop case: rogue text outside of run obj - must fix template
                        paragraph.text = paragraph.text.replace(k, v)

        def iterate_tables(t_list, k, v):
            for table in t_list:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_runs(cell.paragraphs, k, v)

        if key == "pdf_data":
            for paragraph in document.paragraphs:
                if f"<{key}>" in paragraph.text:
                    paragraph.clear()
                    for file in v:
                        self.replace_pdf_extract(paragraph, file["text"], tolerance=2)
                        run = paragraph.add_run()
                        run.add_break()
                        run.add_break()
            return

        if v and key is not None:
            k = "<" + key + ">"
            doc = document
            # some of these are probably unnecessary
            # foot_style = styles['Footer']
            # head_style = styles['Header']
            # t_style = None
            # p_style = None
            run_style = None

            if len(doc.paragraphs) > 0:
                replace_in_runs(doc.paragraphs, k, v)

            if len(doc.tables) > 0:
                iterate_tables(doc.tables, k, v)

            if len(doc.sections) > 0:
                for section in doc.sections:
                    replace_in_runs(section.footer.paragraphs, k, v)
                    iterate_tables(section.footer.tables, k, v)
                    replace_in_runs(section.header.paragraphs, k, v)
                    iterate_tables(section.header.tables, k, v)

    def replace_pdf_extract(self, paragraph, extracted_segments, tolerance=2):
        """
        Given a paragraph and a list of extracted segments from pdf_extract,
        group segments by similar y values (same line) and add them as runs.
        Each segment is a dict with keys "text", "x", "y", "height" and "is_bold".
        """
        # Sort segments first by page then y then x:
        segments = sorted(extracted_segments, key=lambda s: (s["page"], s["y"], s["x"]))
        max_line_width = 510
        
        # Group segments that are on the same line (within a tolerance)
        lines = []
        current_line = []
        current_y = None
        for seg in segments:
            if current_y is None or abs(seg["y"] - current_y) <= tolerance:
                current_line.append(seg)
                current_y = seg["y"]  # use the first segment's y for the line
            else:
                lines.append(current_line)
                current_line = [seg]
                current_y = seg["y"]
        if current_line:
            lines.append(current_line)
        
        # For each group (line), add a run per segment:
        for i, line in enumerate(lines):
            # Sort segments by x position for each line
            line = sorted(line, key=lambda s: s["x"])
            
            for seg in line:
                if seg["text"].startswith("http"):
                    link_text = seg["text"]
                    link_width = seg["x"] + seg["width"]
                    if link_width >= max_line_width or link_text.endswith("-") and next_line is not None:
                         j = i + 1
                         while j < len(lines):
                              next_line = lines[j] if i < len(lines) - 1 else None
                              next_line_width = None
                              if next_line:
                                    next_line_width = sum(s["width"] for s in next_line)
                                    next_line_text = "".join(seg["text"] for seg in next_line)
                              if next_line_width < max_line_width:
                                space_index = next_line_text.find(" ")
                                if space_index != -1:
                                    next_line[0]["text"] = next_line[0]["text"][space_index:]
                                    addition = next_line_text[:space_index]
                                else:
                                    next_line[0]["text"] = ""
                                    addition = next_line_text
                                link_text += addition.strip()
                                break
                              else:
                                link_text += addition.strip()
                                for s in next_line:
                                     s["text"] = ""
                                j += 1
                              
                    html_parser = DocumentHTMLParser(paragraph, self.doc)
                    html_parser.add_hyperlink(paragraph, link_text, link_text, color='5384da')
                else:
                    run = paragraph.add_run(seg["text"])
                    if seg.get("is_bold"):
                        run.bold = True

            # Determine how to join this line with the next (if any)
            if i < len(lines) - 1:
                current_line_text = "".join(seg["text"] for seg in line)
                next_line = lines[i + 1] if i < len(lines) - 1 else None
                if next_line:
                    next_line_text = "".join(seg["text"] for seg in next_line)
                current_line_base = min(seg["y"] for seg in line)
                line_height = max(seg["height"] for seg in line)
                next_line_base = min(seg["y"] for seg in next_line)
                gap = next_line_base - current_line_base
                if gap > 2 * line_height:
                    run = paragraph.add_run()
                    run.add_break()
                    run.add_break()
                elif current_line_text.strip().endswith(":") or re.match(r"^\d+\.", next_line_text):
                    paragraph.add_run().add_break()
                else:
                    paragraph.add_run(" ")


    def insert_image(self, document, k, v, image_path=None, config=None):
        # going to need to write custom logic depending on how images should be placed/styled

        return True

    def insert_custom(self, document, k, v, config=None):
        # perhaps replaces {{custom_object}} with pre-determined text structure with custom style/format

        return True

Alias = str
Mapping = Dict[Alias, Any]
UUIDString = NewType('UUIDString', str)
def validate_uuid_string(uuid_string: str) -> UUIDString:
    try:
        uuid_obj = uuid.UUID(uuid_string)
    except ValueError:
        return False
    return str(uuid_obj) == uuid_string
RelatedResource = NewType("RelatedResource", arches_orm.view_models.node_list.NodeListViewModel)

class GenericTemplateProvider:
    def __init__(self, resource_instance: Resource):
        self.resource_instance = resource_instance
        self.datatype_factory = DataTypeFactory()
        self.config = {}
        # self.tiles not currently used and might need removed
        self.tiles = resource_instance.tiles

    # Legacy, not in use
    def get_value_from_tile(self, tile:Tile, node_id:UUIDString) -> Any:
        """Obtain tile values (Legacy, not in use)

            @param Tile: any tile object

            @param node_id: the target node

            @return Any: tile value or empty string
        """
        current_node = models.Node.objects.get(nodeid=node_id)
        datatype = self.datatype_factory.get_instance(current_node.datatype)
        returnvalue = datatype.get_display_value(tile, current_node)
        return "" if returnvalue is None else returnvalue
    
    # Legacy, not in use
    def get_resource(self, resource_id:UUIDString) -> Resource:
        """ Obtain resource from id

            @param resourceinstanceid: uuid for the resource instance

            @return Resource: an instance of arches.app.models.resource.Resource
        """
        resource = None
        try:
            resource = Resource.objects.filter(pk=resource_id).first()
            
        except Resource.DoesNotExist:
            raise f"Resource ID ({resource_id}) does not exist"
        return resource
    
    def extract(self, node_list) -> Mapping:
        """Drill down a Semantic Node List to obtain actual values

            @param node_list: 
            
            @return Mapping: a mapping with actual values
        """
        children_present = True
        mapping = {}
        while children_present:
            newloop = []
            found_semantic = False
            for item in node_list:
                if isinstance(item[0], str):
                    try:
                        newloop += item[1].items()
                        found_semantic = True
                    except:
                        segment = {item[0] : item[1]}
                        mapping = mapping | self.processDatatypes(segment)
                    node_list = newloop
                    children_present = found_semantic
        return mapping
    
    def merge_mappings(self, original: dict, new: dict):
        merged = original.copy() if original is not None else {}
        for key, value in new.items():
            if 'many_tiles' in self.config and key in self.config['many_tiles'] and key in merged:
                if not isinstance(merged[key], list):
                    merged[key] = [merged[key]]
                merged[key].append(value)
            else:
                merged[key] = value
        return merged

    def extract_from_related_resource(self, prefix:Alias, related_resource:RelatedResource) -> Mapping:
        """Process values from a related resource's node list
        
            @param prefix: the alias for the related resource will be prefixed

            @param related_resource: a node_list from a related resource

            @return processed_items: a partially processed Mapping
        """
        mapping = self.extract(list(related_resource.items()))
        processed_mapping = self.processDatatypes(mapping)
        processed_items = {}
        for processed_item in processed_mapping.items():
            processed_items[f"{prefix}__{processed_item[0]}"] = processed_item[1]
        return processed_items

    def get_mapping(self, config={}) -> Mapping:
        """Creates a mapping for the template engine from the resource instance.

            @param self: must contain resource_instance:Resource

            @param config?: Dict[str:str]

            common config options:

            'include': List[Alias] # if present only specified aliases will have their values mapped

            'exclude': List[Alias] # if present all nodes apart from listed aliases will have their values mapped

            'expand': List[Alias(related resource)] # related resources which should have their values included to the mapping

            'special': Dict[str: List[str]] # The key will be added as an alias, the value determines the mapping value. value[0] determines the function to be used additional strings are paramaters. 

                @example special: {'todays_date': ['make_date', 'today']} # adds 'todays_date' as an alias and uses the make_date function to create today's date
        """
        self.config = config

        # wkrm
        wkrm = get_well_known_resource_model_by_graph_id(self.resource_instance.graph_id)

        # wkri
        resource = wkrm.find(self.resource_instance.resourceinstanceid)

        # Semantic Node List
        semantic_node_list = list(resource.items())

        
        # I overlooked that the keys won't be in the first level of the dict. Would need to flatten the dicts to see if any of the children keys are in the include/exclude lists. Probably not worth doing
        # if "include" in self.config:
            #  semantic_node_list = [item for item in semantic_node_list if item[0] in self.config['include']]
            
        # if "exclude" in self.config:
        #      semantic_node_list = [item for item in semantic_node_list if not item[0] in self.config['exclude']]

        mapping = self.extract(semantic_node_list)  
        if "special" in self.config:
            for special_case in self.config["special"].items():
                if special_case[1] == 'today':
                    mapping[special_case[0]] = datetime.today().strftime("%d/%m/%Y")
                elif special_case[1] == 'user':
                     mapping = self.get_user(mapping, special_case[0])

        return self.processDatatypes(mapping)         

    def processDatatypes(self, mapping:Mapping) -> Mapping:
        """Provides the logic for extracting a node lists's values for different datatypes.

            @param mapping: An unprocessed Mapping

            @return mapping: A processed Mapping with values stringified
        """
        for item in mapping.items():
            alias, value = item

            if isinstance(value, arches_orm.view_models.node_list.NodeListViewModel):
                for node in value:
                    if isinstance(node, arches_orm.view_models.semantic.SemanticViewModel):
                        mapping = self.merge_mappings(mapping, self.extract(list(node.items())))
                # TODO handle node lists that are not semantic e.g bibligraphic source is a resource-instance but has children. Currently we ignore the children 
                mapping[alias] = None
                continue
            if isinstance(value, (arches_orm.view_models.concepts.EmptyConceptValueViewModel, arches_orm.arches_django.datatypes.user.UserViewModel)):
                mapping[alias] = None
                continue
            if isinstance(value, (arches_orm.view_models.resources.RelatedResourceInstanceListViewModel)):             
                if "expand" in self.config and alias in self.config["expand"] and len(value) > 0:
                    for related_resource in value:
                        mapping = mapping | self.extract_from_related_resource(alias, related_resource)
                        mapping[alias] = str(related_resource)
                else:
                    resource_list = []
                    for related_resource in value.data:
                        resource_list.append(str(related_resource))
                    mapping[alias] = resource_list
                continue
            if isinstance(value, (arches_orm.view_models.concepts.ConceptListValueViewModel)):
                concept_list = []
                for datum in value.data:
                    concept_list.append(str(datum))
                mapping[alias] = concept_list
                continue
            if isinstance(value, (arches_orm.view_models.concepts.ConceptValueViewModel)):
                mapping[alias] = str(value)
                continue
            if isinstance(value, (arches_orm.view_models.semantic.SemanticViewModel)):
                dicted_value = {}
                for key in list(value.keys()):
                    if key:
                        dicted_value[key] = value[key]
                mapping = mapping | self.extract(list(dicted_value.items()))
                mapping[alias] = None
                continue
            if isinstance(value, str) and not isinstance(value, (arches_orm.view_models.semantic.SemanticViewModel, arches_orm.view_models.concepts.ConceptValueViewModel, arches_orm.view_models.string.StringViewModel)):
                # Arches ORM doesn't seem to have date datatype
                try:
                    mapping[alias] = datetime.fromisoformat(value).strftime("%d/%m/%Y")
                except Exception as e:
                    pass

                # Arches ORM doesn't seem to have domain-value datatype
                if validate_uuid_string(value):
                     #TODO check only domain-value datatypes are present and make domain logic
                     pass
                continue
                
        return mapping
    
    # not in use but probably a more reusable way of creating special cases
    def make_date(self, operator:str, mapping:Mapping, alias:Alias, vars=None):
         if operator == 'today':
              mapping[alias] = datetime.today().strftime("%d/%m/%Y")
         elif operator == 'delta':
              delta = datetime.timedelta(days=vars.days, months=vars.months, years=vars.years)
              mapping[alias] = datetime.today()+delta.strftime("%d/%m/%Y")
         pass
    
    def get_user(self, mapping:Mapping, alias:Alias) -> Mapping:
        from arches_orm.models import Person
        from arches_orm.adapter import admin
        try:
            with admin():
                person = Person.where(user_account=self.config["user"].id).get()
                person = person[0] if len(person) else self.config["user"] 
                # person = person["name"][0]["full_name"]
                str(person)
        except Resource.DoesNotExist:
            person = self.config["user"]

        mapping[alias] = str(person)
        return mapping
    
    def make_node_list_string(self):
         pass

            
class MonumentTemplateProvider:
    MONUMENT_NAME_NODEGROUP = '676d47f9-9c1c-11ea-9aa0-f875a44e0e11'
    MONUMENT_NAME_NODE = '676d47ff-9c1c-11ea-b07f-f875a44e0e11'

    MONUMENT_SYSTEM_REF_RESOURCE_ID_NODEGROUP = '325a2f2f-efe4-11eb-9b0c-a87eeabdefba'
    MONUMENT_SYSTEM_REF_RESOURCE_ID_NODE = '325a430a-efe4-11eb-810b-a87eeabdefba'

    MONUMENT_CONTACTS_NODEGROUP = 'aa629840-d23e-11ee-9ae7-0242ac180006'
    MONUMENT_CONTACTS_APPLICANT_NODE = 'aa62a736-d23e-11ee-9ae7-0242ac180006'

    MONUMENT_LOCALITIES_ADMIN_AREA_NODEGROUP = '87d38725-f44f-11eb-8d4b-a87eeabdefba'
    MONUMENT_AREA_NAME_NODE = '87d3c3ea-f44f-11eb-b532-a87eeabdefba'

    MONUMENT_ADDRESSES_NODEGROUP = '87d39b25-f44f-11eb-95e5-a87eeabdefba'
    MONUMENT_COUNTY_NODE = '87d3ff32-f44f-11eb-aa82-a87eeabdefba'

    MONUMENT_CM_REFERENCE_NODEGROUP = '3d415e98-d23b-11ee-9373-0242ac180006'
    MONUMENT_CM_REFERENCE_NODE = '3d419020-d23b-11ee-9373-0242ac180006'

    SMC_RECEIVED_DATE_NODEGROUP = 'eeec9986-d23c-11ee-9373-0242ac180006'
    SMC_RECEIVED_DATE_NODE = 'eeec9e68-d23c-11ee-9373-0242ac180006'

    APPLICANT_TITLE_NODEGROUP = '4110f741-1a44-11e9-885e-000d3ab1e588'
    APPLICANT_TITLE_NODE = '6da2f03b-7e55-11ea-8fe5-f875a44e0e11'

    APPLICANT_CONTACT_POINT_NODEGROUP = '2547c12f-9505-11ea-a507-f875a44e0e11'
    APPLICANT_CONTACT_POINT_NODE = '2547c133-9505-11ea-8e49-f875a44e0e11'

    APPLICANT_ADDRESSES_NODEGROUP = '5f93048e-80a9-11ea-b0da-f875a44e0e11'
    APPLICANT_COUNTY_NODE = 'b3a28c1d-effb-11eb-95a1-a87eeabdefba'
    APPLICANT_POSTCODE_NODE = 'b3a27619-effb-11eb-a66d-a87eeabdefba'
    APPLICANT_FULL_ADDRESS_NODE = 'b3a27611-effb-11eb-a79c-a87eeabdefba'

    def __init__(self, resource_instance):
        self.resource_instance = resource_instance
        self.datatype_factory = DataTypeFactory()
        self.tiles = resource_instance.tiles
        self.mapping = {
            "Courtesy Title": "", 
            "Address": "", 
            "County": "", 
            "Postcode": "", 
            "Monument Townland": "", 
            "Monument County": "",
            "Monument Name": "", 
            "SMC": "", # NOT PROVIDED
            "SMR": "", 
            "CM Reference": "", 
            "Recipient Email": "", 
            "Recipient Name": "", 
            "Received Date": "",
            "Send Date": "", # NOT PROVIDED
            "Granted Date": "", # NOT PROVIDED
            "Addendum Condition": "", # NOT PROVIDED
            "Additional Conditions Pretext": "", # NOT PROVIDED
            "Archaeological Inspector": "", # NOT PROVIDED
            "Field Monument Warden": "", # NOT PROVIDED
            "Letter Send Date": "" # NOT PROVIDED
        }

    def get_value_from_tile(self, tile, node_id):
        current_node = models.Node.objects.get(nodeid=node_id)
        datatype = self.datatype_factory.get_instance(current_node.datatype)
        returnvalue = datatype.get_display_value(tile, current_node)
        return "" if returnvalue is None else returnvalue
    
    def get_resource(self, resource_id):

        resource = None
        try:
            resource = Resource.objects.filter(pk=resource_id).first()
        except Resource.DoesNotExist:
            raise f"Resource ID ({resource_id}) does not exist"
        return resource

    def get_mapping(self):

        for tile in self.tiles:
            nodegroup_id = str(tile.nodegroup_id)

            if nodegroup_id == self.MONUMENT_NAME_NODEGROUP:
                self.mapping["Monument Name"] = self.get_value_from_tile(
                    tile, self.MONUMENT_NAME_NODE
                )

            if nodegroup_id == self.MONUMENT_SYSTEM_REF_RESOURCE_ID_NODEGROUP:
                self.mapping["SMR"] = self.get_value_from_tile(
                    tile, self.MONUMENT_SYSTEM_REF_RESOURCE_ID_NODE
                )

            if nodegroup_id == self.MONUMENT_LOCALITIES_ADMIN_AREA_NODEGROUP:
                self.mapping["Monument Townland"] = self.get_value_from_tile(
                    tile, self.MONUMENT_AREA_NAME_NODE
                )
            
            if nodegroup_id == self.MONUMENT_ADDRESSES_NODEGROUP:
                self.mapping["Monument County"] = self.get_value_from_tile(
                    tile, self.MONUMENT_COUNTY_NODE
                )

            if nodegroup_id == self.MONUMENT_CM_REFERENCE_NODEGROUP:
                self.mapping["CM Reference"] = self.get_value_from_tile(
                    tile, self.MONUMENT_CM_REFERENCE_NODE
                )

            if nodegroup_id == self.SMC_RECEIVED_DATE_NODEGROUP:
                self.mapping["Received Date"] = self.get_value_from_tile(
                    tile, self.SMC_RECEIVED_DATE_NODE
                )
            
            if nodegroup_id == self.MONUMENT_CONTACTS_NODEGROUP:
                self.mapping["Recipient Name"] = self.get_value_from_tile(
                    tile, self.MONUMENT_CONTACTS_APPLICANT_NODE
                )

                if not tile.data.get(self.MONUMENT_CONTACTS_APPLICANT_NODE) or not len(tile.data.get(self.MONUMENT_CONTACTS_APPLICANT_NODE)):
                    continue

                applicant_resource_id = tile.data.get(self.MONUMENT_CONTACTS_APPLICANT_NODE)[0].get('resourceId')
                applicant_resource = self.get_resource(applicant_resource_id)
                applicant_resource.load_tiles()
                applicant_tiles = applicant_resource.tiles

                for applicant_tile in applicant_tiles:
                    applicant_nodegroup_id = str(applicant_tile.nodegroup_id)
                    if applicant_nodegroup_id == self.APPLICANT_TITLE_NODEGROUP:
                        self.mapping["Courtesy Title"] = self.get_value_from_tile(
                            applicant_tile, self.APPLICANT_TITLE_NODE
                        )
                    if applicant_nodegroup_id == self.APPLICANT_CONTACT_POINT_NODEGROUP:
                        self.mapping["Recipient Email"] = self.get_value_from_tile(
                            applicant_tile, self.APPLICANT_CONTACT_POINT_NODE
                        )
                    if applicant_nodegroup_id == self.APPLICANT_ADDRESSES_NODEGROUP:
                        self.mapping["County"] = self.get_value_from_tile(
                            applicant_tile, self.APPLICANT_COUNTY_NODE
                        )
                        self.mapping["Postcode"] = self.get_value_from_tile(
                            applicant_tile, self.APPLICANT_POSTCODE_NODE
                        )
                        self.mapping["Address"] = self.get_value_from_tile(
                            applicant_tile, self.APPLICANT_FULL_ADDRESS_NODE
                        )

        return self.mapping
    

class LicenceTemplateProvider:
    LICENCE_NAME_NODEGROUP = '59d65ec0-48b9-11ee-84da-0242ac140007'
    LICENCE_NAME_NODE = '59d6676c-48b9-11ee-84da-0242ac140007'

    LICENCE_SYSTEM_REF_RESOURCE_ID_NODEGROUP = '991c3c74-48b6-11ee-85af-0242ac140007'
    LICENCE_SYSTEM_REF_RESOURCE_ID_NODE = '991c4340-48b6-11ee-85af-0242ac140007'
    LICENCE_NUMBER_NODE = '991c49b2-48b6-11ee-85af-0242ac140007'

    LICENCE_CONTACTS_NODEGROUP = '6397b05c-c443-11ee-94bf-0242ac180006'
    LICENCE_CONTACTS_APPLICANT_NODE = '6d2924b6-5891-11ee-a624-0242ac120004'

    # LICENCE_LOCALITIES_ADMIN_AREA_NODEGROUP = '87d38725-f44f-11eb-8d4b-a87eeabdefba'
    # LICENCE_AREA_NAME_NODE = '87d3c3ea-f44f-11eb-b532-a87eeabdefba'

    # LICENCE_ADDRESSES_NODEGROUP = '87d39b25-f44f-11eb-95e5-a87eeabdefba'
    # LICENCE_COUNTY_NODE = '87d3ff32-f44f-11eb-aa82-a87eeabdefba'

    LICENCE_CM_REFERENCE_NODEGROUP = 'b84fa9c6-bad2-11ee-b3f2-0242ac180006'
    LICENCE_CM_REFERENCE_NODE = 'b84fb182-bad2-11ee-b3f2-0242ac180006'

    REQUEST_DATE_NODEGROUP = '69b2738e-c4d2-11ee-b171-0242ac180006'
    REQUEST_DATE_NODE = 'c6f09242-c4d2-11ee-b171-0242ac180006'

    REPORTED_DATE_NODEGROUP = 'f060583a-6120-11ee-9fd1-0242ac120003'
    REPORTED_DATE_NODE = '0a089af2-dc7a-11ee-8def-0242ac120006'

    REPORT_SUBMITTED_DATE_NODEGROUP = 'f060583a-6120-11ee-9fd1-0242ac120003'    
    REPORT_SUBMITTED_DATE_NODE = '59b77af6-dc6f-11ee-8def-0242ac120006'

    ISSUED_DATE_N1_NODEGROUP = '69b2738e-c4d2-11ee-b171-0242ac180006'
    ISSUED_DATE_N1_NODE = '4129013c-c4d5-11ee-90c5-0242ac180006'
    
    ISSUED_DATE_NODEGROUP = '1887f678-c42d-11ee-bc4b-0242ac180006'
    ISSUED_DATE_NODE = '1887faf6-c42d-11ee-bc4b-0242ac180006'
    
    VALID_UNTIL_DATE_NODEGROUP = '1887f678-c42d-11ee-bc4b-0242ac180006'
    VALID_UNTIL_DATE_NODE = '1887fc86-c42d-11ee-bc4b-0242ac180006'

    DATES_NODEGROUP = '05f6b846-5d49-11ee-911e-0242ac130003'
    PROPOSED_START_DATE_N1_NODE = '84e2983c-5d49-11ee-911e-0242ac130003'
    PROPOSED_START_DATE_NODE = '84e2983c-5d49-11ee-911e-0242ac130003'
    ACKNOWLEDGED_DATE_N1_NODE = 'ed16bb80-5d4a-11ee-9b75-0242ac130003'
    ACTUAL_START_DATE_NODE = '97f6c776-5d4a-11ee-9b75-0242ac130003'
    ACTUAL_END_DATE_NODE = '7ee258e0-5d4a-11ee-9b75-0242ac130003'
    RECIVED_DATE_NODE = '58880bd6-5d4a-11ee-9b75-0242ac130003'
    DURATION_NODE = 'c688af34-d589-11ee-89d9-0242ac180006'

    CMREF_NODE = 'b84fb182-bad2-11ee-b3f2-0242ac180006'

    LPA_PLANNING_REFERENCE_NODEGROUP = '9236156e-bad1-11ee-b3f2-0242ac180006'
    PLANNING_REFERENCE_NODE = '92361d48-bad1-11ee-b3f2-0242ac180006'

    CLASSIFICATION_DATE_NODEGROUP = 'f060583a-6120-11ee-9fd1-0242ac120003'
    CLASSIFICATION_DATE_NODE = 'ea6ea7a8-dc70-11ee-b70c-0242ac120006'
    
    APPLICANT_N1_NODEGROUP = '6397b05c-c443-11ee-94bf-0242ac180006'
    APPLICANT_N1_NODE = 'f870c35e-c447-11ee-8be7-0242ac180006'
    
    APPLICANT_NODEGROUP = '6d290832-5891-11ee-a624-0242ac120004'
    APPLICANT_NODE = '6d2924b6-5891-11ee-a624-0242ac120004'
    COMPANY_NODE = '07d3905c-d58b-11ee-a02f-0242ac180006'
    LICENSEE_NODE = '6d294784-5891-11ee-a624-0242ac120004'
# {'07d3905c-d58b-11ee-a02f-0242ac180006': [{'resourceId': '5a5857b8-7612-48fa-a442-180c2fd22800', 'ontologyProperty': 'ac41d9be-79db-4256-b368-2f4559cfbe55', 'resourceXresourceId': '6e0b30c7-e9ff-4c65-854d-00f4d55fef6e', 'inverseOntologyProperty': 'ac41d9be-79db-4256-b368-2f4559cfbe55'}], '318184a4-d58b-11ee-89d9-0242ac180006': '6a08646f-a2d8-45f7-b2f5-bafa2200e1f8', '4936d1c6-d58b-11ee-a02f-0242ac180006': 'b81d4b16-0633-4d7a-b4b2-5c2d3e2e782e', '6d29144e-5891-11ee-a624-0242ac120004': 'b81d4b16-0633-4d7a-b4b2-5c2d3e2e782e', '6d2924b6-5891-11ee-a624-0242ac120004': [{'resourceId': '76c563e1-c7ce-40c4-86cf-32291b932ebb', 'ontologyProperty': '', 'resourceXresourceId': 'e90cfeb1-2946-4cf1-aed1-7b19a7359b41', 'inverseOntologyProperty': ''}], '6d292772-5891-11ee-a624-0242ac120004': 'b81d4b16-0633-4d7a-b4b2-5c2d3e2e782e', '6d292cf4-5891-11ee-a624-0242ac120004': 'b2489138-dcc5-4151-999e-977809179bb3', '6d293532-5891-11ee-a624-0242ac120004': None, '6d29392e-5891-11ee-a624-0242ac120004': '31a0f492-2c6d-4c1b-8bfa-74da97fa14fc', '6d2941f8-5891-11ee-a624-0242ac120004': 'b81d4b16-0633-4d7a-b4b2-5c2d3e2e782e', '6d2944f0-5891-11ee-a624-0242ac120004': '0a32a247-2b40-4f89-a24e-ce9b88bac793', '6d294784-5891-11ee-a624-0242ac120004': [{'resourceId': 'e12597ad-dfad-4081-bbcb-f7007f50542d', 'ontologyProperty': '', 'resourceXresourceId': 'e3e62a28-2395-46d6-abcf-f68cc18a5436', 'inverseOntologyProperty': ''}]}

    NEW_LICENSEE_NODEGROUP = '6397b05c-c443-11ee-94bf-0242ac180006'
    NEW_LICENSEE_NODE = '66a2157a-c449-11ee-8be7-0242ac180006'
    
    FORMER_LICENSEE_NODEGROUP = '6397b05c-c443-11ee-94bf-0242ac180006'
    FORMER_LICENSEE_NODE = '69fec032-c445-11ee-94bf-0242ac180006'
    
    FILES_NODEGROUP ='8c5356f4-48ce-11ee-8e4e-0242ac140007'
    FILES_NODE ='8c5356f4-48ce-11ee-8e4e-0242ac140007'
    
    ASSOCIATED_ACTIVITIES_NODEGROUP = 'a9f53f00-48b6-11ee-85af-0242ac140007'
    ASSOCIATED_ACTIVITIES_NODE = 'a9f53f00-48b6-11ee-85af-0242ac140007'  

    DECISION_NODEGROUP = '2749ea5a-48cb-11ee-be76-0242ac140007'
    DECISION_NODE = '2749ea5a-48cb-11ee-be76-0242ac140007'    
    
    GRADE_E_NODEGROUP = '69f2eb3c-c430-11ee-94bf-0242ac180006'
    GRADE_E_NODE = '69f30298-c430-11ee-94bf-0242ac180006'
    
    GRADE_E_N1_NODEGROUP = '6397b05c-c443-11ee-94bf-0242ac180006'
    GRADE_E_N1_NODE = '058ccf60-c44d-11ee-94bf-0242ac180006'
    
    GRADE_D_NODEGROUP = 'c9f504b4-c42d-11ee-94bf-0242ac180006'
    GRADE_D_NODE = 'c9f51490-c42d-11ee-94bf-0242ac180006'
    
    GRADE_D_N1_NODEGROUP = '6397b05c-c443-11ee-94bf-0242ac180006'
    GRADE_D_N1_NODE = '6bc892c8-c44d-11ee-94bf-0242ac180006'
    
    GRADE_D_N2_NODEGROUP = '69b2738e-c4d2-11ee-b171-0242ac180006'
    GRADE_D_N2_NODE = '2e7a876e-c4d4-11ee-b171-0242ac180006'
    
    # LICENCE_NUMBER_NODEGROUP = '6de3741e-c502-11ee-86cf-0242ac180006'    
    # LICENCE_NUMBER_NODE = '6de3741e-c502-11ee-86cf-0242ac180006'
    
    PERSON_TITLE_NODEGROUP = '4110f741-1a44-11e9-885e-000d3ab1e588'
    PERSON_TITLE_NODE = '6da2f03b-7e55-11ea-8fe5-f875a44e0e11'
    
    PERSON_CONTACT_POINT_NODEGROUP = '2547c12f-9505-11ea-a507-f875a44e0e11'
    PERSON_CONTACT_POINT_NODE = '2547c133-9505-11ea-8e49-f875a44e0e11'
    
    PERSON_ADDRESSES_NODEGROUP = '5f93048e-80a9-11ea-b0da-f875a44e0e11'
    PERSON_COUNTY_NODE = 'b3a28c1d-effb-11eb-95a1-a87eeabdefba'

    PERSON_CORRESPONDENCE_NODEGROUP = '2547c12f-9505-11ea-a507-f875a44e0e11'
    PERSON_CORRESPONDENCE_NAMES_NODEGROUP = '2beefb51-4084-11eb-9b2b-f875a44e0e11'
    PERSON_CORRESPONDENCE_NAME_NODE = '2beefb56-4084-11eb-bcc5-f875a44e0e11'
    PERSON_CORRESPONDENCE_EMAIL_NODE = '2547c133-9505-11ea-8e49-f875a44e0e11'

    # if nodegroup_id == self.NEW_LICENSEE_NODEGROUP:
    #     self.mapping["New Licensee"] = self.get_value_from_tile(
    #     tile, self.NEW_LICENSEE_NODE
    # )
    PERSON_POSTCODE_NODE = 'b3a27619-effb-11eb-a66d-a87eeabdefba'
    PERSON_FULL_ADDRESS_NODE = 'b3a27611-effb-11eb-a79c-a87eeabdefba'
    # if nodegroup_id == self.NEW_LICENSEE_NODEGROUP:
    #     self.mapping["New Licensee"] = self.get_value_from_tile(
    #     tile, self.NEW_LICENSEE_NODE
    # )


    COMPANY_NAMES_NODEGROUP = 'e8431c5d-8098-11ea-8348-f875a44e0e11'
    COMPANY_NAME_NODE = 'e8431c61-8098-11ea-8b01-f875a44e0e11'

    COMPANY_CORRESPONDENCE_NODEGROUP = '1b6f9cb4-51ae-11eb-a1fe-f875a44e0e11'
    COMPANY_CORRESPONDENCE_NAME_NODE = '1b6f9cb9-51ae-11eb-9ece-f875a44e0e11'
    COMPANY_EMAIL_NODE = '1b6f9cbf-51ae-11eb-b61d-f875a44e0e11'

    COMPANY_ADDRESSES_NODEGROUP = 'af3b0116-29a9-11eb-8333-f875a44e0e11'
    COMPANY_FULL_ADDRESS_NODE = '9e7907c7-eff3-11eb-b606-a87eeabdefba'
    COMPANY_BUILDING_NAME_NODE = '9e7907d3-eff3-11eb-ac11-a87eeabdefba'
    COMPANY_BUILDING_NUMBER_NODE = '9e7907d5-eff3-11eb-a511-a87eeabdefba'
    COMPANY_STREET_NODE = '9e7907d7-eff3-11eb-8e7a-a87eeabdefba'
    COMPANY_LOCALITY_NODE = '9e7907cd-eff3-11eb-b0f1-a87eeabdefba'
    COMPANY_COUNTY_NODE = '9e791cfe-eff3-11eb-9c35-a87eeabdefba'
    COMPANY_POSTCODE_NODE = '9e7907cf-eff3-11eb-8412-a87eeabdefba'

    ACTIVITY_NODEGROUP = 'a9f53f00-48b6-11ee-85af-0242ac140007'
    ACTIVITY_SITE_NAME_NODE = 'a9f53f00-48b6-11ee-85af-0242ac140007'
    ACTIVITY_NODE = 'a9f53f00-48b6-11ee-85af-0242ac140007'

    ACTIVITY_AREA_NODEGROUP = 'a5416b46-f121-11eb-8f2d-a87eeabdefba'
    ACTIVITY_TOWNLAND_NODE = 'a5416b53-f121-11eb-a507-a87eeabdefba'

    ACTIVITY_COUNCIL_NODEGROUP = '5f81a8d4-d7de-11ee-b2c1-0242ac120006'
    ACTIVITY_COUNCIL_NODE = '5f81a8d4-d7de-11ee-b2c1-0242ac120006'

    ACTIVITY_GRID_REFERENCES_NODEGROUP = '33b4430a-16be-11ef-8633-0242ac180006'
    ACTIVITY_IRISH_GRID_REFERENCE_NODE = '4bd349a4-16be-11ef-af79-0242ac180006'

    ACTIVITY_LOCATION_DESCRIPTION_NODEGROUP = 'a541b934-f121-11eb-9d20-a87eeabdefba:'
    ACTIVITY_LOCATION_DESCRIPTION_NODE = 'a5416b40-f121-11eb-9cb6-a87eeabdefba'

    def __init__(self, resource_instance):
        self.resource_instance = resource_instance
        self.datatype_factory = DataTypeFactory()
        self.tiles = resource_instance.tiles
        self.mapping = {
            "Courtesy Title": "", 
            "Address": "", 
            "County": "", 
            "Postcode": "", 
            "Monument Townland": "", 
            "Monument County": "",
            "Monument Name": "", 
            "SMC": "", # NOT PROVIDED
            "SMR": "", 
            "CM Reference": "", 
            "Recipient Email": "", 
            "Recipient Name": "", 
            "Received Date": "",
            "Send Date": "", # NOT PROVIDED
            "Granted Date": "", # NOT PROVIDED
            "Addendum Condition": "", # NOT PROVIDED
            "Additional Conditions Pretext": "", # NOT PROVIDED
            "Archaeological Inspector": "", # NOT PROVIDED
            "Field Monument Warden": "", # NOT PROVIDED
            "Letter Send Date": "" # NOT PROVIDED
        }

    def get_value_from_tile(self, tile, node_id):

        current_node = models.Node.objects.get(nodeid=node_id)
        datatype = self.datatype_factory.get_instance(current_node.datatype)
        returnvalue = datatype.get_display_value(tile, current_node)
        return "" if returnvalue is None else returnvalue
    
    def get_values_from_tile(self, tile):
        values = []
        tileDict = dict(tile.data)
        for node_id in tileDict.keys():     
            current_node = models.Node.objects.get(nodeid=node_id)
            datatype = self.datatype_factory.get_instance(current_node.datatype)
            values.append(datatype.get_display_value(tile, current_node))

        for (idx, value) in enumerate(values):
             if value is None:
                  values[idx] = ""

    
    def get_resource(self, resource_id):

        resource = None
        try:
            resource = Resource.objects.filter(pk=resource_id).first()
        except Resource.DoesNotExist:
            raise f"Resource ID ({resource_id}) does not exist"
        return resource

    def get_mapping(self):

        for tile in self.tiles:
            nodegroup_id = str(tile.nodegroup_id)

            if nodegroup_id == self.COMPANY_ADDRESSES_NODEGROUP:
                        self.mapping["Company County"] = self.get_value_from_tile(
                            tile, self.COMPANY_COUNTY_NODE
                        )
                        self.mapping["Company Postcode"] = self.get_value_from_tile(
                            tile, self.COMPANY_POSTCODE_NODE
                        )
                        self.mapping["Company Building Name"] = self.get_value_from_tile(
                            tile, self.COMPANY_BUILDING_NAME_NODE
                        )
                        self.mapping["Company Building Number"] = self.get_value_from_tile(
                            tile, self.COMPANY_BUILDING_NUMBER_NODE
                        )
                        self.mapping["Company Street"] = self.get_value_from_tile(
                            tile, self.COMPANY_STREET_NODE
                        )
                        self.mapping["Company Locality"] = self.get_value_from_tile(
                            tile, self.COMPANY_LOCALITY_NODE
                        )

                        self.mapping["Company Address"] = self.get_value_from_tile(
                            tile, self.COMPANY_FULL_ADDRESS_NODE
                        )
            if nodegroup_id == self.GRADE_E_NODEGROUP:
                if not tile.data.get(self.GRADE_E_NODE) or not len(tile.data.get(self.GRADE_E_NODE)):
                    continue

                grade_e_resource_id = tile.data.get(self.GRADE_E_NODE)[0].get('resourceId')
                grade_e_resource = self.get_resource(grade_e_resource_id)
                grade_e_resource.load_tiles()
                grade_e_tiles = grade_e_resource.tiles

                for grade_e_tile in grade_e_tiles:

                    grade_e_nodegroup_id = str(grade_e_tile.nodegroup_id)
                    if grade_e_nodegroup_id == self.PERSON_TITLE_NODEGROUP:
                        self.mapping["Grade E Courtesy Title"] = self.get_value_from_tile(
                            grade_e_tile, self.PERSON_TITLE_NODE
                        )
                    if grade_e_nodegroup_id == self.PERSON_CONTACT_POINT_NODEGROUP:
                        self.mapping["Grade E Email"] = self.get_value_from_tile(
                            grade_e_tile, self.PERSON_CONTACT_POINT_NODE
                        )
                    if grade_e_nodegroup_id == self.PERSON_ADDRESSES_NODEGROUP:
                        self.mapping["Grade E County"] = self.get_value_from_tile(
                            grade_e_tile, self.PERSON_COUNTY_NODE
                        )
                        self.mapping["Grade E Postcode"] = self.get_value_from_tile(
                            grade_e_tile, self.PERSON_POSTCODE_NODE
                        )
                        self.mapping["Grade E Address"] = self.get_value_from_tile(
                            grade_e_tile, self.PERSON_FULL_ADDRESS_NODE
                        )
            if nodegroup_id == self.GRADE_D_NODEGROUP:
                if not tile.data.get(self.GRADE_D_NODE) or not len(tile.data.get(self.GRADE_D_NODE)):
                    continue

                grade_d_resource_id = tile.data.get(self.GRADE_D_NODE)[0].get('resourceId')
                grade_d_resource = self.get_resource(grade_d_resource_id)
                grade_d_resource.load_tiles()
                grade_d_tiles = grade_d_resource.tiles

                for grade_d_tile in grade_d_tiles:

                    grade_d_nodegroup_id = str(grade_d_tile.nodegroup_id)
                    if grade_d_nodegroup_id == self.PERSON_TITLE_NODEGROUP:
                        self.mapping["Grade D Courtesy Title"] = self.get_value_from_tile(
                            grade_d_tile, self.PERSON_TITLE_NODE
                        )
                    if grade_d_nodegroup_id == self.PERSON_CONTACT_POINT_NODEGROUP:
                        self.mapping["Grade D Email"] = self.get_value_from_tile(
                            grade_d_tile, self.PERSON_CONTACT_POINT_NODE
                        )
                    if grade_d_nodegroup_id == self.PERSON_ADDRESSES_NODEGROUP:
                        self.mapping["Grade D County"] = self.get_value_from_tile(
                            grade_d_tile, self.PERSON_COUNTY_NODE
                        )
                        self.mapping["Grade D Postcode"] = self.get_value_from_tile(
                            grade_d_tile, self.PERSON_POSTCODE_NODE
                        )
                        self.mapping["Grade D Address"] = self.get_value_from_tile(
                            grade_d_tile, self.PERSON_FULL_ADDRESS_NODE
                        )



            if nodegroup_id == self.LICENCE_NAME_NODEGROUP:
                self.mapping["Licence Name"] = self.get_value_from_tile(
                    tile, self.LICENCE_NAME_NODE
                )


            if nodegroup_id == self.LICENCE_SYSTEM_REF_RESOURCE_ID_NODEGROUP:
                self.mapping["SMR"] = self.get_value_from_tile(
                    tile, self.LICENCE_SYSTEM_REF_RESOURCE_ID_NODE
                )
                self.mapping["Licence Number"] = self.get_value_from_tile(
                    tile, self.LICENCE_NUMBER_NODE
                )


            # if nodegroup_id == self.LICENCE_LOCALITIES_ADMIN_AREA_NODEGROUP:
            #     self.mapping["Licence Townland"] = self.get_value_from_tile(
            #         tile, self.LICENCE_AREA_NAME_NODE
            #     )
            
            # if nodegroup_id == self.LICENCE_ADDRESSES_NODEGROUP:
            #     self.mapping["Licence County"] = self.get_value_from_tile(
            #         tile, self.LICENCE_COUNTY_NODE
            #     )

            if nodegroup_id == self.LICENCE_CM_REFERENCE_NODEGROUP:
                self.mapping["CM Reference"] = self.get_value_from_tile(
                    tile, self.LICENCE_CM_REFERENCE_NODE
                )

            # if nodegroup_id == self.LICENCE_SYSTEM_REF_RESOURCE_ID_NODEGROUP:
            #     self.mapping["Licence Number"] = self.get_value_from_tile(
            #         tile, self.LICENCE_SYSTEM_REF_RESOURCE_ID_NODE
            #     )
            #         tile, self.LICENCE_SYSTEM_REF_RESOURCE_ID_NODE
            #     ))
            # if nodegroup_id == self.SMC_RECEIVED_DATE_NODEGROUP:
            #     self.mapping["Received Date"] = self.get_value_from_tile(
            #         tile, self.SMC_RECEIVED_DATE_NODE
            #     )
            

            if nodegroup_id == self.LICENCE_CONTACTS_NODEGROUP:
                self.mapping["Licence Contacts"] = self.get_value_from_tile(
                    tile, self.LICENCE_CONTACTS_APPLICANT_NODE
                )
            
            if nodegroup_id == self.LICENCE_CM_REFERENCE_NODEGROUP:
                            self.mapping["Licence CM"] = self.get_value_from_tile(
                                tile, self.LICENCE_CM_REFERENCE_NODE
                            )
            
            if nodegroup_id == self.REQUEST_DATE_NODEGROUP:
                            self.mapping["Licence Request Date"] = self.get_value_from_tile(
                                tile, self.REQUEST_DATE_NODE
                            )
            if nodegroup_id == self.REPORTED_DATE_NODEGROUP:
                            self.mapping["Licence Reported Date"] = self.get_value_from_tile(
                                tile, self.REPORTED_DATE_NODE
                            )
            if nodegroup_id == self.DATES_NODEGROUP:
                            self.mapping["Licence Acknowledged Date"] = self.get_value_from_tile(
                                tile, self.ACKNOWLEDGED_DATE_N1_NODE
                            )
                            self.mapping["Proposed Start Date"] = self.get_value_from_tile(
                                tile, self.PROPOSED_START_DATE_N1_NODE
                            )
                            self.mapping["Actual Start Date"] = self.get_value_from_tile(
                                                        tile, self.ACTUAL_START_DATE_NODE
                            )
                            self.mapping["Actual End Date"] = self.get_value_from_tile(
                                                        tile, self.ACTUAL_END_DATE_NODE
                            )

            if nodegroup_id == self.REPORT_SUBMITTED_DATE_NODEGROUP:
                            self.mapping["Licence Submitted Date"] = self.get_value_from_tile(
                                tile, self.REPORT_SUBMITTED_DATE_NODE
                            )
            if nodegroup_id == self.ISSUED_DATE_N1_NODEGROUP:
                        self.mapping["Licence Issued Date n1"] = self.get_value_from_tile(
                            tile, self.ISSUED_DATE_N1_NODE
                        )
            if nodegroup_id == self.ISSUED_DATE_NODEGROUP:
                        self.mapping["Licence Issued Date"] = self.get_value_from_tile(
                            tile, self.ISSUED_DATE_NODE
                        )
            if nodegroup_id == self.VALID_UNTIL_DATE_NODEGROUP:
                        self.mapping["Valid Until"] = self.get_value_from_tile(
                            tile, self.VALID_UNTIL_DATE_NODE
                        )
            

            if nodegroup_id == self.CLASSIFICATION_DATE_NODEGROUP:
                        self.mapping["Classification Date"] = self.get_value_from_tile(
                            tile, self.CLASSIFICATION_DATE_NODE
                        )
            
            if nodegroup_id == self.APPLICANT_N1_NODEGROUP:
                    self.mapping["Applicant n1"] = self.get_value_from_tile(
                        tile, self.APPLICANT_N1_NODE
                    )
            
            if nodegroup_id == self.APPLICANT_NODEGROUP:
                    self.mapping["Applicant"] = self.get_value_from_tile(
                    tile, self.APPLICANT_NODE
                )
            
            if nodegroup_id == self.NEW_LICENSEE_NODEGROUP:
                    self.mapping["New Licensee"] = self.get_value_from_tile(
                    tile, self.NEW_LICENSEE_NODE
                )
            
            if nodegroup_id == self.FILES_NODEGROUP:
                    self.mapping["Files"] = self.get_value_from_tile(
                    tile, self.FILES_NODE
                )
            
            
            if nodegroup_id == self.ASSOCIATED_ACTIVITIES_NODEGROUP:
                    self.mapping["Associated Activities"] = self.get_value_from_tile(
                    tile, self.ASSOCIATED_ACTIVITIES_NODE
                )
            
            if nodegroup_id == self.GRADE_E_NODEGROUP:
                    self.mapping["Grade E"] = self.get_value_from_tile(
                    tile, self.GRADE_E_NODE
                )
            
            if nodegroup_id == self.GRADE_E_N1_NODEGROUP:
                    self.mapping["Grade E n1"] = self.get_value_from_tile(
                    tile, self.GRADE_E_N1_NODE
                )
            
            if nodegroup_id == self.GRADE_D_NODEGROUP:
                    self.mapping["Grade D"] = self.get_value_from_tile(
                    tile, self.GRADE_D_NODE
                )
            
            if nodegroup_id == self.GRADE_D_N1_NODEGROUP:
                    self.mapping["Grade D n1"] = self.get_value_from_tile(
                    tile, self.GRADE_D_N1_NODE
                )
            if nodegroup_id == self.GRADE_D_N2_NODEGROUP:
                    self.mapping["Grade D n2"] = self.get_value_from_tile(
                    tile, self.GRADE_D_N2_NODE
                )
                    
            # if nodegroup_id == self.LICENCE_NUMBER_NODEGROUP:
            #         self.mapping["Licence Number"] = self.get_value_from_tile(
            #         tile, self.LICENCE_NUMBER_NODE
            #     )
            
            # if nodegroup_id == self.PERSON_TITLE_NODEGROUP:
            #         self.mapping["Applicant Title"] = self.get_value_from_tile(
            #         tile, self.PERSON_TITLE_NODE
            #     )
            
            # if nodegroup_id == self.PERSON_CONTACT_POINT_NODEGROUP:
            #         self.mapping["Applicant Contact Point"] = self.get_value_from_tile(
            #         tile, self.PERSON_CONTACT_POINT_NODE
            #     )

            if nodegroup_id == self.APPLICANT_NODEGROUP:
                self.mapping["Recipient Name"] = self.get_value_from_tile(
                    tile, self.APPLICANT_NODE
                )

                #     tile, self.LICENCE_CONTACTS_APPLICANT_NODE
                # ))

                if not tile.data.get(self.LICENCE_CONTACTS_APPLICANT_NODE) or not len(tile.data.get(self.LICENCE_CONTACTS_APPLICANT_NODE)):
                    continue

                applicant_resource_id = tile.data.get(self.LICENCE_CONTACTS_APPLICANT_NODE)[0].get('resourceId')
                applicant_resource = self.get_resource(applicant_resource_id)
                applicant_resource.load_tiles()
                applicant_tiles = applicant_resource.tiles

                if not tile.data.get(self.COMPANY_NODE) or not len(tile.data.get(self.COMPANY_NODE)):
                    continue

                # {'07d3905c-d58b-11ee-a02f-0242ac180006': [{'resourceId': '5a5857b8-7612-48fa-a442-180c2fd22800', 'ontologyProperty': 'ac41d9be-79db-4256-b368-2f4559cfbe55', 'resourceXresourceId': '6e0b30c7-e9ff-4c65-854d-00f4d55fef6e', 'inverseOntologyProperty': 'ac41d9be-79db-4256-b368-2f4559cfbe55'}], '318184a4-d58b-11ee-89d9-0242ac180006': '6a08646f-a2d8-45f7-b2f5-bafa2200e1f8', '4936d1c6-d58b-11ee-a02f-0242ac180006': 'b81d4b16-0633-4d7a-b4b2-5c2d3e2e782e', '6d29144e-5891-11ee-a624-0242ac120004': 'b81d4b16-0633-4d7a-b4b2-5c2d3e2e782e', '6d2924b6-5891-11ee-a624-0242ac120004': [{'resourceId': '76c563e1-c7ce-40c4-86cf-32291b932ebb', 'ontologyProperty': '', 'resourceXresourceId': 'e90cfeb1-2946-4cf1-aed1-7b19a7359b41', 'inverseOntologyProperty': ''}], '6d292772-5891-11ee-a624-0242ac120004': 'b81d4b16-0633-4d7a-b4b2-5c2d3e2e782e', '6d292cf4-5891-11ee-a624-0242ac120004': 'b2489138-dcc5-4151-999e-977809179bb3', '6d293532-5891-11ee-a624-0242ac120004': None, '6d29392e-5891-11ee-a624-0242ac120004': '31a0f492-2c6d-4c1b-8bfa-74da97fa14fc', '6d2941f8-5891-11ee-a624-0242ac120004': 'b81d4b16-0633-4d7a-b4b2-5c2d3e2e782e', '6d2944f0-5891-11ee-a624-0242ac120004': '0a32a247-2b40-4f89-a24e-ce9b88bac793', '6d294784-5891-11ee-a624-0242ac120004': [{'resourceId': 'e12597ad-dfad-4081-bbcb-f7007f50542d', 'ontologyProperty': '', 'resourceXresourceId': 'e3e62a28-2395-46d6-abcf-f68cc18a5436', 'inverseOntologyProperty': ''}]}


                company_resource_id = tile.data.get(self.COMPANY_NODE)[0].get('resourceId')
                company_resource = self.get_resource(company_resource_id)
                company_resource.load_tiles()
                company_tiles = company_resource.tiles


                for applicant_tile in applicant_tiles:

                    applicant_nodegroup_id = str(applicant_tile.nodegroup_id)
                    if applicant_nodegroup_id == self.PERSON_TITLE_NODEGROUP:
                        self.mapping["Applicant Courtesy Title"] = self.get_value_from_tile(
                            applicant_tile, self.PERSON_TITLE_NODE
                        )
                    if applicant_nodegroup_id == self.PERSON_CONTACT_POINT_NODEGROUP:
                        self.mapping["Applicant Email"] = self.get_value_from_tile(
                            applicant_tile, self.PERSON_CONTACT_POINT_NODE
                        )

                    if applicant_nodegroup_id == self.PERSON_ADDRESSES_NODEGROUP:
                        self.mapping["Applicant County"] = self.get_value_from_tile(
                            applicant_tile, self.PERSON_COUNTY_NODE
                        )
                        self.mapping["Applicant Postcode"] = self.get_value_from_tile(
                            applicant_tile, self.PERSON_POSTCODE_NODE
                        )
                        self.mapping["Applicant Address"] = self.get_value_from_tile(
                            applicant_tile, self.PERSON_FULL_ADDRESS_NODE
                        )
                    if applicant_nodegroup_id == self.PERSON_TITLE_NODEGROUP:
                        self.mapping["Applicant Courtesy Title"] = self.get_value_from_tile(
                            applicant_tile, self.PERSON_TITLE_NODE
                        )
                    if applicant_nodegroup_id == self.PERSON_CONTACT_POINT_NODEGROUP:
                        self.mapping["Recipient Email"] = self.get_value_from_tile(
                            applicant_tile, self.PERSON_CONTACT_POINT_NODE
                        )

                    if applicant_nodegroup_id == self.PERSON_ADDRESSES_NODEGROUP:
                        self.mapping["Recipient County"] = self.get_value_from_tile(
                            applicant_tile, self.PERSON_COUNTY_NODE
                        )
                        self.mapping["Recipient Postcode"] = self.get_value_from_tile(
                            applicant_tile, self.PERSON_POSTCODE_NODE
                        )
                        self.mapping["Recipient Address"] = self.get_value_from_tile(
                            applicant_tile, self.PERSON_FULL_ADDRESS_NODE
                        )
                        

# PERSON_CORRESPONDENCE_NODEGROUP = '2547c12f-9505-11ea-a507-f875a44e0e11'
# PERSON_CORRESPONDENCE_NAMES_NODEGROUP = '2beefb51-4084-11eb-9b2b-f875a44e0e11'
# PERSON_CORRESPONDENCE_NAME_NODE = '2beefb56-4084-11eb-bcc5-f875a44e0e11'
# PERSON_CORRESPONDENCE_EMAIL_NODE = '2547c133-9505-11ea-8e49-f875a44e0e11'

                for company_tile in company_tiles:

                    company_nodegroup_id = str(company_tile.nodegroup_id)
                    if company_nodegroup_id == self.COMPANY_NAMES_NODEGROUP:
                        self.mapping["Company Name"] = self.get_value_from_tile(
                            company_tile, self.COMPANY_NAME_NODE
                        )
                    if company_nodegroup_id == self.COMPANY_CORRESPONDENCE_NODEGROUP:
                        self.mapping["Company Correspondence Name"] = self.get_value_from_tile(
                            company_tile, self.COMPANY_CORRESPONDENCE_NAME_NODE
                        )
                    if company_nodegroup_id == self.COMPANY_CORRESPONDENCE_NODEGROUP:
                        self.mapping["Company Email"] = self.get_value_from_tile(
                            company_tile, self.COMPANY_EMAIL_NODE
                        )
                    if company_nodegroup_id == self.COMPANY_ADDRESSES_NODEGROUP:
                        self.mapping["Company County"] = self.get_value_from_tile(
                            company_tile, self.COMPANY_COUNTY_NODE
                        )
                        self.mapping["Company Postcode"] = self.get_value_from_tile(
                            company_tile, self.COMPANY_POSTCODE_NODE
                        )
                        self.mapping["Company Building Name"] = self.get_value_from_tile(
                            company_tile, self.COMPANY_BUILDING_NAME_NODE
                        )
                        self.mapping["Company Building Number"] = self.get_value_from_tile(
                            company_tile, self.COMPANY_BUILDING_NUMBER_NODE
                        )
                        self.mapping["Company Street"] = self.get_value_from_tile(
                            company_tile, self.COMPANY_STREET_NODE
                        )
                        self.mapping["Company Locality"] = self.get_value_from_tile(
                            company_tile, self.COMPANY_LOCALITY_NODE
                        )

                        self.mapping["Company Address"] = self.get_value_from_tile(
                            company_tile, self.COMPANY_FULL_ADDRESS_NODE
                        )
                


            if nodegroup_id == self.DECISION_NODEGROUP:
                self.mapping["DECISION"] = self.get_value_from_tile(
                   tile, self.DECISION_NODE
                )


            if nodegroup_id == self.GRADE_D_NODEGROUP:
                self.get_values_from_tile(tile)
                self.mapping["Grade D"] = self.get_value_from_tile(
                    tile, self.GRADE_D_NODE
                )

            if nodegroup_id == self.GRADE_E_NODEGROUP:
                self.mapping["Grade E Name"] = self.get_value_from_tile(
                    tile, self.GRADE_E_NODE
                )

                if not tile.data.get(self.GRADE_E_NODE) or not len(tile.data.get(self.GRADE_E_NODE)):
                    continue

                grade_e_resource_id = tile.data.get(self.GRADE_E_NODE)[0].get('resourceId')
                grade_e_resource = self.get_resource(grade_e_resource_id)
                grade_e_resource.load_tiles()
                grade_e_tiles = grade_e_resource.tiles

                for grade_e_tile in grade_e_tiles:

                    grade_e_nodegroup_id = str(grade_e_tile.nodegroup_id)
                    if grade_e_nodegroup_id == self.PERSON_TITLE_NODEGROUP:
                        self.mapping["Grade E Courtesy Title"] = self.get_value_from_tile(
                            grade_e_tile, self.PERSON_TITLE_NODE
                        )
                    if grade_e_nodegroup_id == self.PERSON_CONTACT_POINT_NODEGROUP:
                        self.mapping["Grade E Email"] = self.get_value_from_tile(
                            grade_e_tile, self.PERSON_CONTACT_POINT_NODE
                        )
                    if grade_e_nodegroup_id == self.PERSON_ADDRESSES_NODEGROUP:
                        self.mapping["Grade E County"] = self.get_value_from_tile(
                            grade_e_tile, self.PERSON_COUNTY_NODE
                        )
                        self.mapping["Grade E Postcode"] = self.get_value_from_tile(
                            grade_e_tile, self.PERSON_POSTCODE_NODE
                        )
                        self.mapping["Grade E Address"] = self.get_value_from_tile(
                            grade_e_tile, self.PERSON_FULL_ADDRESS_NODE
                        )

            if nodegroup_id == self.ACTIVITY_NODEGROUP:
                self.mapping["Site Name"] = self.get_value_from_tile(
                    tile, self.ACTIVITY_SITE_NAME_NODE
                )
                activity_resource_id = tile.data.get(self.ACTIVITY_NODE)[0].get('resourceId')
                activity_resource = self.get_resource(activity_resource_id)
                activity_resource.load_tiles()
                activity_tiles = activity_resource.tiles
                for activity_tile in activity_tiles:
                    activity_nodegroup_id = str(activity_tile.nodegroup_id)  
                    if activity_nodegroup_id == self.ACTIVITY_AREA_NODEGROUP:
                       self.mapping["Townland"] = self.get_value_from_tile(
                           activity_tile, self.ACTIVITY_TOWNLAND_NODE
                       )
                    if activity_nodegroup_id == self.ACTIVITY_COUNCIL_NODEGROUP:
                       self.mapping["Council"] = self.get_value_from_tile(
                           activity_tile, self.ACTIVITY_COUNCIL_NODE
                       )
                    if activity_nodegroup_id == self.ACTIVITY_GRID_REFERENCES_NODEGROUP:
                       self.mapping["Irish Grid Reference"] = self.get_value_from_tile(
                           activity_tile, self.ACTIVITY_IRISH_GRID_REFERENCE_NODE
                       )
                    if activity_nodegroup_id == self.ACTIVITY_LOCATION_DESCRIPTION_NODEGROUP:
                       self.mapping["Location Description"] = self.get_value_from_tile(
                           activity_tile, self.ACTIVITY_LOCATION_DESCRIPTION_NODE
                       )
        return self.mapping

class DocumentHTMLParser(HTMLParser):
    def __init__(self, paragraph, document):

        HTMLParser.__init__(self)
        self.document = document
        self.paragraph = paragraph
        self.table = None
        self.table_cols = 0
        self.table_rows = 0
        self.max_cols_reached = False
        self.td_cursor = False
        self.hyperlink = False
        self.list_style = "ul"
        self.ol_counter = 1
        self.run = self.paragraph.add_run()

    def insert_paragraph_after(self, paragraph, text=None, style=None):
        """Insert a new paragraph after the given paragraph."""
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        return new_para

    def add_hyperlink(self, paragraph, url, text, color=None, underline=None):

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = self.paragraph.part
        r_id = part.relate_to(
            url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(
            docx.oxml.shared.qn("r:id"),
            r_id,
        )

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement("w:r")

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement("w:rPr")

        # Add color if it is given
        if not color is None:
            c = docx.oxml.shared.OxmlElement("w:color")
            c.set(docx.oxml.shared.qn("w:val"), color)
            rPr.append(c)  # #5384da ; rgb(83,132,218)

        # Remove underlining if it is requested
        if not underline:
            u = docx.oxml.shared.OxmlElement("w:u")
            u.set(docx.oxml.shared.qn("w:val"), "none")
            rPr.append(u)

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

        return hyperlink

    def insert_into_paragraph_and_feed(self, html):

        html = html.replace("\n\n", "<br>")
        self.run = self.paragraph.add_run()
        self.feed(html)

    def handle_starttag(self, tag, attrs):

        self.run = self.paragraph.add_run()
        if tag == "i" or tag == "em":
            self.run.italic = True
        if tag == "b" or tag == "strong":
            self.run.bold = True
        if tag == "s":
            self.run.strike = True
        if tag == "u":
            self.run.underline = True
        if tag == "ol":
            self.list_style = "ol"
        if tag == "ul":
            self.list_style = "ul"
        if tag in ["br", "ul", "ol"]:
            self.run.add_break()
        if tag == "li":
            if self.list_style == "ul":
                self.run.add_text("● ")
            else:
                self.run.add_text(str(self.ol_counter) + ". ")
                self.ol_counter += 1
        if tag == "p":
            self.run.add_break()
            # self.run.add_break()
            # self.run.add_tab()
        if tag == "a":
            self.hyperlink = attrs[0][1]
        if tag == "table":
            self.table = self.document.add_table(self.table_rows, self.table_cols)
            self.table.autofit = True
        if tag == "tr":
            self.table_rows += 1
            self.table.add_row()
        if tag == "td":
            self.table_cols += 1
            if self.max_cols_reached is False:
                self.table.add_column(1)
            self.td_cursor = True

    def handle_endtag(self, tag):
        
        if tag in ["br", "li", "ul", "ol"]:
            self.run.add_break()
        self.run = self.paragraph.add_run()
        if tag == "ol":
            self.ol_counter = 1
        if tag == "table":
            tbl = self.table._tbl
            p = self.paragraph._p
            p.addnext(tbl)
            self.table = None
            self.table_cols = 0
            self.table_rows = 0
        if tag == "tr":
            self.table_cols = 0
            self.max_cols_reached = True
        if tag == "td":
            self.td_cursor = False

    def handle_data(self, data):

        if "&#39;" in data:
            data = data.replace("&#39;", "'")

        if self.hyperlink is not False:
            blue = docx.shared.RGBColor(83, 132, 218)
            color = blue.__str__()
            self.add_hyperlink(self.paragraph, self.hyperlink, data, color)
            self.hyperlink = False
        elif self.td_cursor is True:
            self.table.cell(self.table_rows - 1, self.table_cols - 1).add_paragraph(
                data
            )  # formatting?
        else:
            self.run.add_text(data)

    def handle_entityref(self, name):
        c = chr(name2codepoint[name])
        self.run.add_text(c)

    def handle_charref(self, name):
        if name.startswith("x"):
            c = chr(int(name[1:], 16))
        else:
            c = chr(int(name))
        self.run.add_text(c)
